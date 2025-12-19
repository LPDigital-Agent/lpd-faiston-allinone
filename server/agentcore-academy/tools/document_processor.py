# =============================================================================
# Document Processor Tool - PDF/DOCX/TXT Extraction
# =============================================================================
# Extracts text content from user-uploaded documents for training creation.
#
# Supported formats:
# - PDF: Uses PyPDF2 for text extraction (fallback to OCR if needed)
# - DOCX: Uses python-docx for structured extraction
# - TXT/CSV/JSON: Direct text reading with encoding detection
#
# Usage:
#   text = extract_text_from_document(file_bytes, "report.pdf", "application/pdf")
#
# S3 Integration:
#   documents = process_documents_from_s3(training_id)
# =============================================================================

import io
import json
import csv
import os
from typing import Optional, Dict, Any, List, Tuple
from pathlib import Path

import boto3
from botocore.exceptions import ClientError

# Lazy imports for cold start optimization
_pypdf2 = None
_docx = None


def _get_pypdf2():
    """Lazy load PyPDF2."""
    global _pypdf2
    if _pypdf2 is None:
        from PyPDF2 import PdfReader
        _pypdf2 = PdfReader
    return _pypdf2


def _get_docx():
    """Lazy load python-docx."""
    global _docx
    if _docx is None:
        import docx
        _docx = docx
    return _docx


# =============================================================================
# Configuration
# =============================================================================

# Maximum file size (10 MB)
MAX_FILE_SIZE = 10 * 1024 * 1024

# Maximum text length to extract (500K characters)
MAX_TEXT_LENGTH = 500_000

# Supported MIME types
SUPPORTED_MIME_TYPES = {
    "application/pdf": "pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
    "application/msword": "doc",  # Legacy .doc (limited support)
    "text/plain": "txt",
    "text/csv": "csv",
    "application/json": "json",
    "text/markdown": "md",
}

# S3 Configuration
S3_BUCKET = os.getenv("TRAININGS_BUCKET", "hive-academy-trainings-prod")
AWS_REGION = os.getenv("AWS_REGION", "us-east-2")


# =============================================================================
# Text Extraction Functions
# =============================================================================


def extract_text_from_pdf(file_bytes: bytes, filename: str = "document.pdf") -> str:
    """
    Extract text from PDF file.

    Args:
        file_bytes: PDF file content as bytes
        filename: Original filename (for logging)

    Returns:
        Extracted text content

    Raises:
        ValueError: If PDF is encrypted or unreadable
    """
    PdfReader = _get_pypdf2()

    try:
        pdf_file = io.BytesIO(file_bytes)
        reader = PdfReader(pdf_file)

        # Check if encrypted
        if reader.is_encrypted:
            raise ValueError(f"PDF '{filename}' is encrypted and cannot be processed")

        # Extract text from all pages
        text_parts = []
        for i, page in enumerate(reader.pages):
            try:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(f"[Pagina {i + 1}]\n{page_text}")
            except Exception as e:
                print(f"[DocumentProcessor] Warning: Failed to extract page {i + 1}: {e}")

        if not text_parts:
            raise ValueError(f"PDF '{filename}' contains no extractable text")

        return "\n\n".join(text_parts)

    except Exception as e:
        if "encrypted" in str(e).lower():
            raise ValueError(f"PDF '{filename}' is encrypted: {e}")
        raise ValueError(f"Failed to process PDF '{filename}': {e}")


def extract_text_from_docx(file_bytes: bytes, filename: str = "document.docx") -> str:
    """
    Extract text from DOCX file.

    Args:
        file_bytes: DOCX file content as bytes
        filename: Original filename (for logging)

    Returns:
        Extracted text content (paragraphs and tables)

    Raises:
        ValueError: If DOCX is corrupted or unreadable
    """
    docx = _get_docx()

    try:
        docx_file = io.BytesIO(file_bytes)
        doc = docx.Document(docx_file)

        text_parts = []

        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)

        # Extract tables
        for table in doc.tables:
            table_text = []
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    table_text.append(row_text)
            if table_text:
                text_parts.append("[Tabela]\n" + "\n".join(table_text))

        if not text_parts:
            raise ValueError(f"DOCX '{filename}' contains no extractable text")

        return "\n\n".join(text_parts)

    except Exception as e:
        raise ValueError(f"Failed to process DOCX '{filename}': {e}")


def extract_text_from_txt(file_bytes: bytes, filename: str = "document.txt") -> str:
    """
    Extract text from TXT/Markdown file.

    Args:
        file_bytes: Text file content as bytes
        filename: Original filename (for logging)

    Returns:
        Text content

    Raises:
        ValueError: If file cannot be decoded
    """
    # Try common encodings
    encodings = ["utf-8", "latin-1", "cp1252", "iso-8859-1"]

    for encoding in encodings:
        try:
            text = file_bytes.decode(encoding)
            if text.strip():
                return text
        except (UnicodeDecodeError, LookupError):
            continue

    raise ValueError(f"Failed to decode text file '{filename}' - unsupported encoding")


def extract_text_from_csv(file_bytes: bytes, filename: str = "document.csv") -> str:
    """
    Extract text from CSV file (converts to readable format).

    Args:
        file_bytes: CSV file content as bytes
        filename: Original filename (for logging)

    Returns:
        Text representation of CSV data

    Raises:
        ValueError: If CSV cannot be parsed
    """
    # Decode file
    text = extract_text_from_txt(file_bytes, filename)

    try:
        reader = csv.reader(io.StringIO(text))
        rows = list(reader)

        if not rows:
            raise ValueError(f"CSV '{filename}' is empty")

        # Convert to readable format
        text_parts = []
        headers = rows[0] if rows else []

        for i, row in enumerate(rows):
            if i == 0:
                text_parts.append("Cabecalho: " + " | ".join(headers))
            else:
                row_items = [f"{headers[j]}: {cell}" for j, cell in enumerate(row) if j < len(headers) and cell.strip()]
                if row_items:
                    text_parts.append(f"Linha {i}: " + ", ".join(row_items))

        return "\n".join(text_parts)

    except csv.Error as e:
        raise ValueError(f"Failed to parse CSV '{filename}': {e}")


def extract_text_from_json(file_bytes: bytes, filename: str = "document.json") -> str:
    """
    Extract text from JSON file (converts to readable format).

    Args:
        file_bytes: JSON file content as bytes
        filename: Original filename (for logging)

    Returns:
        Text representation of JSON data

    Raises:
        ValueError: If JSON cannot be parsed
    """
    text = extract_text_from_txt(file_bytes, filename)

    try:
        data = json.loads(text)

        # Convert to readable format
        def flatten_json(obj, prefix="") -> List[str]:
            items = []
            if isinstance(obj, dict):
                for key, value in obj.items():
                    new_key = f"{prefix}.{key}" if prefix else key
                    items.extend(flatten_json(value, new_key))
            elif isinstance(obj, list):
                for i, item in enumerate(obj):
                    items.extend(flatten_json(item, f"{prefix}[{i}]"))
            else:
                items.append(f"{prefix}: {obj}")
            return items

        flattened = flatten_json(data)
        return "\n".join(flattened)

    except json.JSONDecodeError as e:
        raise ValueError(f"Failed to parse JSON '{filename}': {e}")


# =============================================================================
# Main Extraction Function
# =============================================================================


def extract_text_from_document(
    file_bytes: bytes,
    filename: str,
    mime_type: Optional[str] = None,
) -> Dict[str, Any]:
    """
    Extract text from a document file.

    Args:
        file_bytes: File content as bytes
        filename: Original filename
        mime_type: MIME type (auto-detected if not provided)

    Returns:
        Dict with:
        - text: Extracted text content
        - char_count: Number of characters
        - format: Detected format
        - filename: Original filename

    Raises:
        ValueError: If file type is unsupported or extraction fails
    """
    # Validate file size
    if len(file_bytes) > MAX_FILE_SIZE:
        raise ValueError(f"File '{filename}' exceeds maximum size of {MAX_FILE_SIZE // 1024 // 1024}MB")

    # Determine format from extension if mime_type not provided
    ext = Path(filename).suffix.lower().lstrip(".")

    if mime_type and mime_type in SUPPORTED_MIME_TYPES:
        file_format = SUPPORTED_MIME_TYPES[mime_type]
    elif ext in ("pdf", "docx", "doc", "txt", "csv", "json", "md"):
        file_format = ext
    else:
        raise ValueError(f"Unsupported file type: {filename} (mime: {mime_type})")

    # Extract text based on format
    if file_format == "pdf":
        text = extract_text_from_pdf(file_bytes, filename)
    elif file_format in ("docx", "doc"):
        text = extract_text_from_docx(file_bytes, filename)
    elif file_format in ("txt", "md"):
        text = extract_text_from_txt(file_bytes, filename)
    elif file_format == "csv":
        text = extract_text_from_csv(file_bytes, filename)
    elif file_format == "json":
        text = extract_text_from_json(file_bytes, filename)
    else:
        raise ValueError(f"Unsupported format: {file_format}")

    # Truncate if too long
    if len(text) > MAX_TEXT_LENGTH:
        text = text[:MAX_TEXT_LENGTH] + f"\n\n[Truncado: texto excedeu {MAX_TEXT_LENGTH} caracteres]"
        print(f"[DocumentProcessor] Warning: Text truncated for '{filename}'")

    return {
        "text": text,
        "char_count": len(text),
        "format": file_format,
        "filename": filename,
    }


# =============================================================================
# S3 Integration
# =============================================================================


def get_s3_client():
    """Get configured S3 client."""
    return boto3.client("s3", region_name=AWS_REGION)


def download_document_from_s3(
    s3_key: str,
    bucket: Optional[str] = None,
) -> Tuple[bytes, str]:
    """
    Download a document from S3.

    Args:
        s3_key: S3 object key
        bucket: S3 bucket (defaults to trainings bucket)

    Returns:
        Tuple of (file_bytes, content_type)

    Raises:
        ValueError: If download fails
    """
    bucket = bucket or S3_BUCKET
    s3 = get_s3_client()

    try:
        response = s3.get_object(Bucket=bucket, Key=s3_key)
        file_bytes = response["Body"].read()
        content_type = response.get("ContentType", "application/octet-stream")
        return file_bytes, content_type

    except ClientError as e:
        error_code = e.response.get("Error", {}).get("Code", "Unknown")
        if error_code == "NoSuchKey":
            raise ValueError(f"Document not found in S3: {s3_key}")
        raise ValueError(f"Failed to download from S3: {e}")


def process_documents_from_s3(
    training_id: str,
    document_keys: List[str],
    bucket: Optional[str] = None,
) -> List[Dict[str, Any]]:
    """
    Process multiple documents from S3.

    Args:
        training_id: Training ID (for logging)
        document_keys: List of S3 keys to process
        bucket: S3 bucket (defaults to trainings bucket)

    Returns:
        List of extraction results with status

    Example:
        results = process_documents_from_s3(
            training_id="abc123",
            document_keys=[
                "trainings/abc123/documents/1/report.pdf",
                "trainings/abc123/documents/2/data.csv",
            ]
        )
    """
    results = []

    for s3_key in document_keys:
        filename = Path(s3_key).name

        try:
            # Download from S3
            file_bytes, content_type = download_document_from_s3(s3_key, bucket)

            # Extract text
            extraction = extract_text_from_document(file_bytes, filename, content_type)

            results.append({
                "s3_key": s3_key,
                "status": "success",
                **extraction,
            })

            print(f"[DocumentProcessor] Processed '{filename}': {extraction['char_count']} chars")

        except Exception as e:
            print(f"[DocumentProcessor] Error processing '{filename}': {e}")
            results.append({
                "s3_key": s3_key,
                "filename": filename,
                "status": "error",
                "error": str(e),
            })

    return results


def upload_document_to_s3(
    file_bytes: bytes,
    training_id: str,
    filename: str,
    version: int = 1,
    bucket: Optional[str] = None,
) -> str:
    """
    Upload a document to S3.

    Args:
        file_bytes: File content
        training_id: Training ID
        filename: Original filename
        version: Document version (for versioning)
        bucket: S3 bucket (defaults to trainings bucket)

    Returns:
        S3 key of uploaded document

    Raises:
        ValueError: If upload fails
    """
    bucket = bucket or S3_BUCKET
    s3 = get_s3_client()

    # Sanitize filename
    safe_filename = "".join(c for c in filename if c.isalnum() or c in "._-")
    s3_key = f"trainings/{training_id}/documents/{version}/{safe_filename}"

    # Determine content type
    ext = Path(filename).suffix.lower()
    content_types = {
        ".pdf": "application/pdf",
        ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ".doc": "application/msword",
        ".txt": "text/plain",
        ".csv": "text/csv",
        ".json": "application/json",
        ".md": "text/markdown",
    }
    content_type = content_types.get(ext, "application/octet-stream")

    try:
        s3.put_object(
            Bucket=bucket,
            Key=s3_key,
            Body=file_bytes,
            ContentType=content_type,
            Metadata={
                "training_id": training_id,
                "original_filename": filename,
                "version": str(version),
            },
        )

        print(f"[DocumentProcessor] Uploaded '{filename}' to s3://{bucket}/{s3_key}")
        return s3_key

    except ClientError as e:
        raise ValueError(f"Failed to upload to S3: {e}")


def generate_presigned_upload_url(
    training_id: str,
    filename: str,
    version: int = 1,
    expires_in: int = 3600,
    bucket: Optional[str] = None,
) -> Dict[str, str]:
    """
    Generate a pre-signed URL for direct browser upload.

    Args:
        training_id: Training ID
        filename: Target filename
        version: Document version
        expires_in: URL expiration in seconds (default 1 hour)
        bucket: S3 bucket (defaults to trainings bucket)

    Returns:
        Dict with:
        - upload_url: Pre-signed PUT URL
        - s3_key: S3 key where file will be stored

    Raises:
        ValueError: If URL generation fails
    """
    bucket = bucket or S3_BUCKET
    s3 = get_s3_client()

    # Sanitize filename
    safe_filename = "".join(c for c in filename if c.isalnum() or c in "._-")
    s3_key = f"trainings/{training_id}/documents/{version}/{safe_filename}"

    try:
        upload_url = s3.generate_presigned_url(
            "put_object",
            Params={
                "Bucket": bucket,
                "Key": s3_key,
            },
            ExpiresIn=expires_in,
        )

        return {
            "upload_url": upload_url,
            "s3_key": s3_key,
        }

    except ClientError as e:
        raise ValueError(f"Failed to generate upload URL: {e}")


# =============================================================================
# Utility Functions
# =============================================================================


def get_supported_formats() -> List[Dict[str, str]]:
    """
    Return list of supported document formats for UI.

    Returns:
        List of format descriptions
    """
    return [
        {"extension": "pdf", "name": "PDF", "mime": "application/pdf"},
        {"extension": "docx", "name": "Word Document", "mime": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"},
        {"extension": "txt", "name": "Text File", "mime": "text/plain"},
        {"extension": "csv", "name": "CSV Spreadsheet", "mime": "text/csv"},
        {"extension": "json", "name": "JSON Data", "mime": "application/json"},
        {"extension": "md", "name": "Markdown", "mime": "text/markdown"},
    ]


def validate_file_type(filename: str, mime_type: Optional[str] = None) -> bool:
    """
    Check if a file type is supported.

    Args:
        filename: File name with extension
        mime_type: Optional MIME type

    Returns:
        True if supported, False otherwise
    """
    ext = Path(filename).suffix.lower().lstrip(".")
    supported_exts = {"pdf", "docx", "doc", "txt", "csv", "json", "md"}

    if ext in supported_exts:
        return True

    if mime_type and mime_type in SUPPORTED_MIME_TYPES:
        return True

    return False
